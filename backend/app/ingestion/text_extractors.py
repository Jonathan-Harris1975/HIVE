from __future__ import annotations

import csv
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

SUPPORTED_TEXT_SUFFIXES = {
    ".txt", ".md", ".log", ".py", ".js", ".ts", ".tsx", ".jsx", ".css", ".html", ".htm",
    ".json", ".csv", ".pdf", ".docx", ".xlsx",
}


@dataclass(frozen=True)
class ExtractedText:
    text: str
    supported: bool
    extractor: str
    suffix: str
    truncated: bool = False
    metadata: dict[str, object] = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return len(self.text)

    def as_dict(self) -> dict[str, object]:
        return {
            "supported": self.supported,
            "extractor": self.extractor,
            "suffix": self.suffix,
            "char_count": self.char_count,
            "truncated": self.truncated,
            "metadata": self.metadata,
        }


def extract_text(path: Path, mime_type: str | None = None, max_chars: int | None = None) -> str:  # noqa: ARG001
    return extract_text_with_metadata(path, mime_type=mime_type, max_chars=max_chars).text


def extract_text_with_metadata(
    path: Path,
    mime_type: str | None = None,  # noqa: ARG001
    *,
    max_chars: int | None = None,
    pdf_max_pages: int = 40,
    csv_max_rows: int = 2000,
    xlsx_max_rows_per_sheet: int = 500,
    xlsx_max_sheets: int = 12,
    docx_max_table_rows: int = 2000,
) -> ExtractedText:
    suffix = path.suffix.lower()
    limit = _safe_limit(max_chars)

    try:
        if suffix in {".txt", ".md", ".log", ".py", ".js", ".ts", ".tsx", ".jsx", ".css"}:
            text = path.read_text(encoding="utf-8", errors="replace")
            text, truncated = _truncate(text, limit)
            return ExtractedText(text=text, supported=True, extractor="plain_text", suffix=suffix, truncated=truncated)
        if suffix in {".html", ".htm"}:
            return extract_html(path, max_chars=limit)
        if suffix == ".json":
            text = json.dumps(json.loads(path.read_text(encoding="utf-8")), indent=2, ensure_ascii=False)
            text, truncated = _truncate(text, limit)
            return ExtractedText(text=text, supported=True, extractor="json", suffix=suffix, truncated=truncated)
        if suffix == ".csv":
            return extract_csv(path, limit_rows=csv_max_rows, max_chars=limit)
        if suffix == ".pdf":
            return extract_pdf(path, max_pages=pdf_max_pages, max_chars=limit)
        if suffix == ".docx":
            return extract_docx(path, max_table_rows=docx_max_table_rows, max_chars=limit)
        if suffix == ".xlsx":
            return extract_xlsx(path, max_rows_per_sheet=xlsx_max_rows_per_sheet, max_sheets=xlsx_max_sheets, max_chars=limit)
    except Exception as exc:  # extraction should report, not crash ingestion
        return ExtractedText(
            text="",
            supported=True,
            extractor=_extractor_name(suffix),
            suffix=suffix,
            truncated=False,
            metadata={"error": str(exc)},
        )

    return ExtractedText(text="", supported=False, extractor="unsupported", suffix=suffix)


def extract_html(path: Path, max_chars: int | None = None) -> ExtractedText:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        text, truncated = _truncate(raw, max_chars)
        return ExtractedText(text=text, supported=True, extractor="html_raw", suffix=path.suffix.lower(), truncated=truncated)
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text("\n")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    text, truncated = _truncate(text, max_chars)
    return ExtractedText(text=text, supported=True, extractor="html", suffix=path.suffix.lower(), truncated=truncated)


def extract_csv(path: Path, limit_rows: int = 2000, max_chars: int | None = None) -> ExtractedText:
    rows: list[str] = []
    truncated = False
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle)
        for index, row in enumerate(reader):
            if index >= limit_rows:
                rows.append("[TRUNCATED_ROWS]")
                truncated = True
                break
            rows.append(", ".join(row))
    text, char_truncated = _truncate("\n".join(rows), max_chars)
    return ExtractedText(
        text=text,
        supported=True,
        extractor="csv",
        suffix=".csv",
        truncated=truncated or char_truncated,
        metadata={"rows_read_limit": limit_rows},
    )


def extract_pdf(path: Path, max_pages: int = 40, max_chars: int | None = None) -> ExtractedText:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ExtractedText(text="", supported=True, extractor="pdf", suffix=".pdf", metadata={"error": "pypdf not installed"})
    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    pages_to_read = min(page_count, max(1, max_pages))
    parts: list[str] = []
    truncated = page_count > pages_to_read
    for index in range(pages_to_read):
        page_text = reader.pages[index].extract_text() or ""
        if page_text.strip():
            parts.append(f"# Page {index + 1}\n{page_text}")
        if max_chars and sum(len(part) for part in parts) >= max_chars:
            truncated = True
            break
    text, char_truncated = _truncate("\n\n".join(parts), max_chars)
    return ExtractedText(
        text=text,
        supported=True,
        extractor="pdf",
        suffix=".pdf",
        truncated=truncated or char_truncated,
        metadata={"pages_total": page_count, "pages_read": pages_to_read, "pages_read_limit": max_pages},
    )


def extract_docx(path: Path, max_table_rows: int = 2000, max_chars: int | None = None) -> ExtractedText:
    try:
        from docx import Document
    except ImportError:
        return ExtractedText(text="", supported=True, extractor="docx", suffix=".docx", metadata={"error": "python-docx not installed"})
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables: list[str] = []
    table_rows_seen = 0
    table_rows_truncated = False
    for table in document.tables:
        for row in table.rows:
            if table_rows_seen >= max_table_rows:
                table_rows_truncated = True
                break
            tables.append(" | ".join(cell.text.strip() for cell in row.cells))
            table_rows_seen += 1
        if table_rows_truncated:
            break
    text, char_truncated = _truncate("\n".join(paragraphs + tables), max_chars)
    return ExtractedText(
        text=text,
        supported=True,
        extractor="docx",
        suffix=".docx",
        truncated=table_rows_truncated or char_truncated,
        metadata={
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
            "table_rows_read": table_rows_seen,
            "table_rows_read_limit": max_table_rows,
        },
    )


def extract_xlsx(path: Path, max_rows_per_sheet: int = 500, max_sheets: int = 12, max_chars: int | None = None) -> ExtractedText:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ExtractedText(text="", supported=True, extractor="xlsx", suffix=".xlsx", metadata={"error": "openpyxl not installed"})
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    sheets = workbook.worksheets[:max_sheets]
    truncated = len(workbook.worksheets) > len(sheets)
    rows_read = 0
    for sheet in sheets:
        parts.append(f"# Sheet: {sheet.title}")
        for index, row in enumerate(sheet.iter_rows(values_only=True)):
            if index >= max_rows_per_sheet:
                parts.append("[TRUNCATED_ROWS]")
                truncated = True
                break
            rows_read += 1
            parts.append(" | ".join("" if value is None else str(value) for value in row))
        if max_chars and sum(len(part) for part in parts) >= max_chars:
            truncated = True
            break
    workbook.close()
    text, char_truncated = _truncate("\n".join(parts), max_chars)
    return ExtractedText(
        text=text,
        supported=True,
        extractor="xlsx",
        suffix=".xlsx",
        truncated=truncated or char_truncated,
        metadata={
            "sheet_count": len(workbook.worksheets),
            "sheets_read": len(sheets),
            "sheets_read_limit": max_sheets,
            "rows_read": rows_read,
            "rows_per_sheet_limit": max_rows_per_sheet,
        },
    )


def chunk_text(text: str, max_chars: int = 6000, overlap: int = 400) -> Iterable[str]:
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = max(0, end - overlap)
    return chunks


def _truncate(text: str, max_chars: int | None) -> tuple[str, bool]:
    if not max_chars or max_chars <= 0 or len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def _safe_limit(max_chars: int | None) -> int | None:
    if max_chars is None:
        return None
    try:
        value = int(max_chars)
    except Exception:
        return None
    return max(1, value)


def _extractor_name(suffix: str) -> str:
    return suffix.lstrip(".") or "unknown"
