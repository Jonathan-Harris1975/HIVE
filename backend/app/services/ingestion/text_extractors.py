from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable


def extract_text(path: Path, mime_type: str | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".log", ".py", ".js", ".ts", ".tsx", ".jsx", ".css", ".html", ".htm"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".json":
        return json.dumps(json.loads(path.read_text(encoding="utf-8")), indent=2, ensure_ascii=False)
    if suffix == ".csv":
        return extract_csv(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".xlsx":
        return extract_xlsx(path)
    return ""


def extract_csv(path: Path, limit_rows: int = 2000) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle)
        for index, row in enumerate(reader):
            if index >= limit_rows:
                rows.append("[TRUNCATED]")
                break
            rows.append(", ".join(row))
    return "\n".join(rows)


def extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables: list[str] = []
    for table in document.tables:
        for row in table.rows:
            tables.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs + tables)


def extract_xlsx(path: Path, max_rows_per_sheet: int = 500) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""
    workbook = load_workbook(path, read_only=False, data_only=False)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for index, row in enumerate(sheet.iter_rows(values_only=True)):
            if index >= max_rows_per_sheet:
                parts.append("[TRUNCATED]")
                break
            parts.append(" | ".join("" if value is None else str(value) for value in row))
    return "\n".join(parts)


def chunk_text(text: str, max_chars: int = 6000, overlap: int = 400) -> Iterable[str]:
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = max(0, end - overlap)
    return chunks
